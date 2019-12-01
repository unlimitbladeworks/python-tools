# -*- coding: utf-8 -*-
"""
@Author  : Sy
@File    : word_demo.py
@Time    : 2019-12-01 10:27
@desc    : word 简单操作
"""

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

# 创建文档对象
document = Document()
document.styles['Normal'].font.name = '宋体'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 添加标题
document.add_heading('第一个标题头', level=0)
document.add_heading('第一个标题头', level=1)
document.add_heading('第一个标题头', level=2)

# 文字段落
p = document.add_paragraph('我是第一段落起始语句!')
# 加粗
p.add_run('加粗的咪哥杂谈').bold = True
# 修改字体大小
p.add_run('普普通通的咪哥杂谈').font.size = Pt(24)
# 斜体
p.add_run('斜体的咪哥杂谈').italic = True
# 新起一段落
p2 = document.add_paragraph('我是新起一行的段落!')

document.add_paragraph(
    '无序列表1', style='List Bullet'
)
document.add_paragraph(
    '无序列表2', style='List Bullet'
)
document.add_paragraph(
    '有序列表1', style='List Number'
)
document.add_paragraph(
    '有序列表2', style='List Number'
)

from docx.enum.text import WD_ALIGN_PARAGRAPH

demo = (
    ('1', '小明', '男'),
    ('2', '小花', '女'),
    ('3', '小二', '男')
)

table = document.add_table(rows=1, cols=3)
tr_cells = table.rows[0].cells
# 表格头文字居中
tr_cells[0].add_paragraph(text='序号').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr_cells[1].add_paragraph(text='姓名').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr_cells[2].add_paragraph(text='性别').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 遍历demo中的内容，进行填充
for id, name, gendar in demo:
    row_cells = table.add_row().cells
    row_cells[0].add_paragraph(text=id).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_cells[1].add_paragraph(text=name).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    row_cells[2].add_paragraph(text=gendar).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 分页
document.add_page_break()

# 保存 word 名字叫 myword
document.save('myword4.docx')
